import re
from pathlib import Path
from typing import Optional

from docx import Document
from pypdf import PdfReader


class ResumeParsingError(Exception):
    """Raised when a resume file cannot be parsed properly."""


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from a PDF resume file.

    Args:
        file_path: Path to the PDF file.

    Returns:
        Extracted text as a string.

    Raises:
        ResumeParsingError: If file cannot be read or extracted text is empty.
    """
    try:
        reader = PdfReader(file_path)
        text_parts = []

        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        text = "\n".join(text_parts).strip()

        if not text:
            raise ResumeParsingError("No readable text could be extracted from the PDF.")

        return text

    except Exception as e:
        raise ResumeParsingError(f"Failed to parse PDF resume: {e}") from e


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from a DOCX resume file.

    Args:
        file_path: Path to the DOCX file.

    Returns:
        Extracted text as a string.

    Raises:
        ResumeParsingError: If file cannot be read or extracted text is empty.
    """
    try:
        doc = Document(file_path)
        text_parts = []

        for para in doc.paragraphs:
            if para.text and para.text.strip():
                text_parts.append(para.text.strip())

        # Also read basic table content if resume uses tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_parts.append(" | ".join(row_text))

        text = "\n".join(text_parts).strip()

        if not text:
            raise ResumeParsingError("No readable text could be extracted from the DOCX.")

        return text

    except Exception as e:
        raise ResumeParsingError(f"Failed to parse DOCX resume: {e}") from e


def clean_resume_text(text: str) -> str:
    """
    Clean extracted resume text.

    Args:
        text: Raw extracted text.

    Returns:
        Cleaned text.
    """
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[^\S\n]+", " ", text)
    return text.strip()


def extract_resume_text(file_path: str) -> str:
    """
    Extract and clean text from a resume file.
    Supports PDF and DOCX.

    Args:
        file_path: Path to the uploaded resume.

    Returns:
        Cleaned resume text.

    Raises:
        ResumeParsingError: If format unsupported or parsing fails.
    """
    path = Path(file_path)

    if not path.exists():
        raise ResumeParsingError(f"File not found: {file_path}")

    suffix = path.suffix.lower()

    if suffix == ".pdf":
        text = extract_text_from_pdf(file_path)
    elif suffix == ".docx":
        text = extract_text_from_docx(file_path)
    else:
        raise ResumeParsingError(
            f"Unsupported resume format: {suffix}. Please upload a PDF or DOCX file."
        )

    cleaned = clean_resume_text(text)

    if len(cleaned) < 100:
        raise ResumeParsingError(
            "The resume text looks too short or unreadable after extraction. "
            "Please upload a clearer PDF or DOCX file."
        )

    return cleaned